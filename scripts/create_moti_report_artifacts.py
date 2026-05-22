from __future__ import annotations

import base64
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
SCREENSHOTS = REPORTS / "screenshots"
TODAY = date.today().isoformat()

GREEN = "58CC02"
BLUE = "1CB0F6"
RED = "FF4B4B"
YELLOW = "FFC800"
INK = "1F2937"
MUTED = "667085"
LIGHT = "F4F7FA"
BORDER = "DDE5ED"


PDFS = [
    Path("F:/1. 프로젝트 실무 개요.pdf"),
    Path("F:/2. 프로젝트실무 진행 절차1.pdf"),
    Path("F:/3. 프로젝트 계획서 작성 방법.pdf"),
    Path("F:/4. 프로젝트 계획서 작성 방법2.pdf"),
    Path("F:/5. 프로젝트 설계1.pdf"),
    Path("F:/6. 프로젝트 설계2.pdf"),
]


FEATURES = [
    ("AUTH-01", "기본 관리자 로그인", "아이디 admin, 비밀번호 admin으로 즉시 로그인할 수 있는 기본 계정을 제공한다.", "완료", "로그인 화면"),
    ("AUTH-02", "회원가입", "신규 사용자가 아이디, 이름, 비밀번호를 등록할 수 있으며 기존 로그인 흐름과 병행된다.", "완료", "회원가입 화면"),
    ("AUTH-03", "아이디 찾기", "등록 정보 기반으로 아이디를 확인하는 보조 화면을 제공한다.", "완료", "아이디 찾기 화면"),
    ("AUTH-04", "비밀번호 찾기", "비밀번호 확인/재설정 안내를 위한 보조 화면을 제공한다.", "완료", "비밀번호 찾기 화면"),
    ("AUTH-05", "로그아웃", "사이드바 로그아웃 버튼을 통해 확인 다이얼로그 후 로그인 화면으로 이동한다.", "완료", "사이드바, 다이얼로그"),
    ("BRAND-01", "앱명 Moti 적용", "화면, 패키지명, 설치 프로그램 표시명, 스플래시 등에 Moti 이름을 적용한다.", "완료", "전역"),
    ("BRAND-02", "Moti 아이콘 적용", "로그인, 앱 아이콘, favicon, 설치/제거 아이콘을 M 중심 아이콘으로 통일한다.", "완료", "public, assets"),
    ("INSTALL-01", "설치 동의문 한글 표시", "NSIS 설치 개인정보/카메라 동의문 인코딩 문제를 해결하고 Moti 문구로 갱신한다.", "완료", "privacy-agreement"),
    ("UI-01", "커스텀 알림/확인창", "브라우저 기본 alert/confirm 대신 앱 디자인에 맞는 모달 다이얼로그를 사용한다.", "완료", "AppDialog"),
    ("UI-02", "버튼/텍스트 대비 개선", "흰 배경에서 보이지 않던 취소, 회원가입, AI 피드백, 로그아웃 관련 텍스트 색상을 개선한다.", "완료", "Button, Login, Statistics"),
    ("UI-03", "로그아웃 위험 색상", "로그아웃 버튼과 확인 다이얼로그에 노란색 대신 부드러운 빨간색 위험 톤을 적용한다.", "완료", "Sidebar, AppDialog"),
    ("DASH-01", "대시보드", "학습 모드 카드와 최근 상태를 한 화면에서 확인하고 학습 화면으로 진입한다.", "완료", "Dashboard"),
    ("MODE-01", "거북목 모드", "목/어깨 기준점으로 고개 전방 이동과 목 각도를 분석하는 모드를 제공한다.", "완료", "LearningSession"),
    ("MODE-02", "키보드 모드", "키보드 사용 자세 점검용 모드를 제공하며 확장 가능한 모드 데이터 구조로 관리한다.", "완료", "modes.tsx"),
    ("MODE-03", "어깨 모드", "좌우 어깨 균형과 상체 기울기 점검용 모드를 제공한다.", "완료", "modes.tsx"),
    ("MODE-04", "안구 모드", "화면 거리/눈 피로 점검 콘셉트의 학습 모드를 제공한다.", "완료", "modes.tsx"),
    ("CAM-01", "웹캠 미리보기", "학습 화면에서 웹캠 화면을 크게 표시하고 사용자가 자세를 쉽게 확인하도록 한다.", "완료", "PostureMonitor"),
    ("CAM-02", "1280x720 요청", "우선 1280x720 exact로 요청하고, 장치가 지원하지 않으면 ideal 1280x720로 fallback한다.", "완료", "useWebcam"),
    ("CAM-03", "웹캠 화면 꽉 채우기", "카메라 화면을 object-cover로 표시하여 좌우 검은 여백을 최소화한다.", "완료", "PostureMonitor"),
    ("CAM-04", "카메라 상태 표시 위치", "카메라 미리보기/AI 분석 상태를 웹캠 확인 화면 우측에 표시하고 하단 시간 표시를 제거한다.", "완료", "LearningSession"),
    ("AI-01", "MediaPipe Pose 로딩", "CDN 기반 MediaPipe Pose 모델과 drawing utils를 로딩해 자세 랜드마크를 계산한다.", "완료", "useMediaPipe"),
    ("AI-02", "관절 표시", "웹캠 화면 위 canvas에 어깨, 귀, 관절점과 연결선을 표시한다.", "완료", "PostureMonitor"),
    ("AI-03", "자세 점수 계산", "프레임별 자세 상태를 기반으로 점수와 경고 상태를 산출한다.", "완료", "postureCalculator, PostureMonitor"),
    ("AI-04", "시작/중지 제어", "자세교정 시작/중지 버튼으로 분석 루프와 화면 상태를 제어한다.", "완료", "LearningSession"),
    ("AI-05", "트레이 이동 안내", "트레이 버튼을 누르면 실제 앱 동작을 설명하는 커스텀 안내 다이얼로그를 표시한다.", "완료", "LearningSession"),
    ("STAT-01", "통계 요약 카드", "오늘 평균 점수, 개선도, 위험 경고 등 핵심 지표를 카드로 제공한다.", "완료", "Statistics"),
    ("STAT-02", "자세 점수 변화 차트", "시간대별 자세 점수 변화를 그래프로 확인한다.", "완료", "Statistics/Recharts"),
    ("STAT-03", "질환 위험도 분석", "거북목, 손목터널, 안구건조증, 어깨 비대칭 위험도를 시각화한다.", "완료", "Statistics"),
    ("STAT-04", "AI 맞춤 피드백", "누적 자세 데이터를 해석한 추천 스트레칭과 작업 환경 개선 제안을 제공한다.", "완료", "Statistics"),
    ("HIST-01", "일간 학습이력", "선택한 날짜의 학습 세션과 점수, 시간, 모드 정보를 표시한다.", "완료", "LearningHistory"),
    ("HIST-02", "주간 학습이력", "일~토 주간 캘린더를 표시하고 이전/다음 주 이동을 지원한다.", "완료", "LearningHistory"),
    ("HIST-03", "월간 학습이력", "월 단위 캘린더와 날짜별 학습 마커를 표시한다.", "완료", "LearningHistory"),
    ("HIST-04", "기간 이동/오늘", "일/주/월 화면에서 이전/다음 기간 이동과 오늘로 돌아가기 버튼을 제공한다.", "완료", "LearningHistory"),
    ("HIST-05", "주말 색상", "토요일은 파란색, 일요일은 빨간색으로 표시하여 달력 가독성을 높인다.", "완료", "LearningHistory"),
    ("SET-01", "설정 화면", "사용자 설정, 알림, 테마, 개인정보/통계 초기화 등의 항목을 제공한다.", "완료", "Settings"),
    ("SET-02", "통계 삭제 확인", "통계 삭제 등 위험 작업은 커스텀 다이얼로그로 재확인한다.", "완료", "Settings"),
    ("ROUTE-01", "HashRouter 라우팅", "Electron 환경에서 안정적으로 동작하도록 HashRouter 기반 라우팅을 사용한다.", "완료", "App"),
    ("DATA-01", "로컬 인증 저장소", "데모 단계에서 localStorage 기반 사용자/세션 저장소로 인증 흐름을 유지한다.", "완료", "authStore"),
    ("QA-01", "타입체크", "빌드 대신 TypeScript 프로젝트 타입체크(tsc -b)를 기본 검증 절차로 사용한다.", "완료", "tsconfig"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(9)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.2
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Title", 24, GREEN),
        ("Heading 1", 16, INK),
        ("Heading 2", 13, INK),
        ("Heading 3", 11, INK),
    ]:
        st = styles[style_name]
        st.font.name = "Malgun Gothic"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)


def add_cover(doc: Document, title: str, subtitle: str, doc_type: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Moti")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(GREEN)

    title_p = doc.add_paragraph(style="Title")
    title_p.add_run(title)
    subtitle_p = doc.add_paragraph()
    subtitle_p.add_run(subtitle).font.size = Pt(12)
    subtitle_p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for cell in row.cells:
            set_cell_shading(cell, LIGHT)
    meta = [
        ("문서 유형", doc_type),
        ("프로젝트", "AI 기반 VDT 증후군 자세 교정 데스크톱 앱"),
        ("작성일", TODAY),
        ("기술 스택", "Electron, React, TypeScript, MediaPipe Pose, Recharts, localStorage"),
    ]
    for i, (k, v) in enumerate(meta):
        set_cell_text(table.cell(i, 0), k, True)
        set_cell_text(table.cell(i, 1), v)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[str]], widths: list[float] | None = None) -> None:
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if widths:
        for i, w in enumerate(widths):
            table.columns[i].width = Inches(w)
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, GREEN)
        set_cell_text(cell, h, True)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], str(value))


def extract_pdf_outline() -> list[str]:
    snippets: list[str] = []
    for pdf in PDFS:
        if not pdf.exists():
            continue
        try:
            reader = PdfReader(str(pdf))
            text = "\n".join(page.extract_text() or "" for page in reader.pages[:2])
            clean = " ".join(text.split())
            snippets.append(f"{pdf.name}: {clean[:180]}")
        except Exception as exc:
            snippets.append(f"{pdf.name}: 읽기 실패({exc})")
    return snippets


def create_feature_spec() -> None:
    doc = Document()
    set_doc_styles(doc)
    add_cover(doc, "Moti 기능 명세서", "현재 구현된 기능을 화면·모듈·검증 관점에서 정리한 상세 명세", "기능 명세서")
    doc.add_heading("1. 개요", level=1)
    doc.add_paragraph("Moti는 PC 웹캠과 AI 자세 분석을 활용해 VDT 증후군 위험 자세를 학습·점검하는 데스크톱 애플리케이션이다. 본 문서는 현재 구현된 모든 주요 기능을 기능 단위로 분해해 설명한다.")
    doc.add_heading("2. 기능 목록", level=1)
    add_table(doc, ["ID", "기능", "상세 설명", "상태", "관련 화면/모듈"], FEATURES, [0.75, 1.3, 3.15, 0.6, 1.4])
    doc.add_heading("3. 화면별 기능 매핑", level=1)
    mapping = [
        ("로그인", "관리자 기본 로그인, 회원가입 이동, 아이디/비밀번호 찾기 이동, Moti 아이콘 표시"),
        ("회원가입", "신규 사용자 등록, 로그인 화면 복귀"),
        ("대시보드", "학습 모드 선택, 최근 상태/학습 진입"),
        ("자세 학습", "웹캠 미리보기, 1280x720 요청, MediaPipe 분석, 관절 표시, 점수/상태 표시, 시작/중지, 트레이 안내"),
        ("통계", "평균 점수, 개선도, 위험 경고, 점수 변화 차트, 질환 위험도, AI 피드백"),
        ("학습 이력", "일간/주간/월간 캘린더, 이전/다음 기간 이동, 오늘 복귀, 토/일 색상 구분, 날짜별 세션 표시"),
        ("설정", "알림/테마/개인정보/통계 초기화, 로그아웃 및 위험 작업 확인"),
    ]
    add_table(doc, ["화면", "주요 기능"], mapping, [1.4, 5.1])
    doc.add_heading("4. 검증 기준", level=1)
    for item in [
        "TypeScript 타입체크(tsc -b)가 통과해야 한다.",
        "로그인 admin/admin으로 주요 화면에 접근할 수 있어야 한다.",
        "웹캠이 1280x720 exact를 지원하지 않아도 ideal fallback으로 미리보기가 열려야 한다.",
        "학습 화면은 웹캠 영역이 우선 보이고, 하단 시간초/스크롤 없이 핵심 정보가 보여야 한다.",
        "학습 이력은 일/주/月 전환과 이전/다음/오늘 이동이 가능해야 한다.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.save(REPORTS / "Moti_기능명세서.docx")


def create_final_report() -> None:
    doc = Document()
    set_doc_styles(doc)
    add_cover(doc, "Moti 프로젝트 최종 보고서", "AI 기반 VDT 증후군 자세 교정 프로그램 구현 결과 보고", "최종 보고서")
    doc.add_heading("초록", level=1)
    doc.add_paragraph("Moti는 웹캠 기반 자세 인식과 데스크톱 UX를 결합해 사용자가 거북목, 어깨 불균형, 키보드 사용 자세, 안구 피로 등 VDT 증후군 관련 위험 요인을 학습하고 관리할 수 있도록 설계한 애플리케이션이다. 본 프로젝트는 Electron/React 기반 클라이언트, MediaPipe Pose 기반 자세 분석, 학습 이력 및 통계 UI, 사용자 친화적인 알림/설정 흐름을 통합하는 것을 목표로 진행되었다.")
    doc.add_heading("1. 서론", level=1)
    doc.add_heading("1.1 개발 배경", level=2)
    doc.add_paragraph("장시간 PC 사용 환경에서는 목, 어깨, 손목, 눈 피로가 누적되기 쉽다. 기존 알림형 도구는 사용자의 실제 자세를 반영하기 어렵기 때문에, 웹캠과 AI 분석을 활용한 실시간 자세 피드백이 필요하다.")
    doc.add_heading("1.2 개발 목적", level=2)
    for item in [
        "웹캠 기반 자세 분석으로 사용자의 자세 문제를 즉시 인지시킨다.",
        "학습 모드별 점수, 경고, 이력, 통계를 제공해 반복 학습을 유도한다.",
        "설치형 데스크톱 앱 수준의 자연스러운 화면 흐름과 알림 UI를 제공한다.",
        "보고서, 기능 명세, 화면 흐름 자료까지 포함해 시연과 발표가 가능한 프로젝트 산출물을 구성한다.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("2. 참고한 보고서 작성 기준", level=1)
    doc.add_paragraph("제공된 프로젝트 실무/계획서/설계 PDF 예시를 참고하여 서론, 요구사항, 설계, 구현, 테스트, 결론 순서로 최종 보고서를 구성하였다. PDF 예시에서 확인한 공통 구조는 다음과 같다.")
    for s in extract_pdf_outline():
        doc.add_paragraph(s, style="List Bullet")
    doc.add_heading("3. 요구사항 및 서비스 범위", level=1)
    add_table(doc, ["구분", "내용"], [
        ("인증", "로그인, 회원가입, 아이디/비밀번호 찾기, 관리자 데모 계정"),
        ("학습", "거북목/키보드/어깨/안구 모드, 웹캠 미리보기, 분석 시작/중지"),
        ("AI 분석", "MediaPipe Pose 랜드마크 추출, 목 각도 계산, canvas 관절 표시"),
        ("통계", "점수 변화, 위험도, AI 피드백, 추천 스트레칭"),
        ("이력", "일간/주간/월간 캘린더, 날짜별 학습 내역, 주말 색상 구분"),
        ("설정", "테마, 알림, 통계 삭제, 로그아웃 확인, 커스텀 다이얼로그"),
    ], [1.2, 5.3])
    doc.add_heading("4. 시스템 설계", level=1)
    doc.add_paragraph("Moti는 Electron Shell 위에서 React Renderer를 구동하는 구조이며, 카메라 스트림은 브라우저 WebRTC API를 통해 획득한다. MediaPipe Pose는 CDN 스크립트 로딩 후 비디오 프레임을 처리하고, 분석 결과는 canvas 오버레이와 점수 UI로 표시된다.")
    add_table(doc, ["계층", "구성 요소", "역할"], [
        ("Desktop Shell", "Electron main/preload", "앱 창, 스플래시, 패키징, 설치 아이콘, 트레이 안내 기반"),
        ("Renderer", "React, HashRouter", "로그인, 대시보드, 학습, 통계, 이력, 설정 화면"),
        ("Vision", "useWebcam, useMediaPipe", "카메라 스트림 획득, 1280x720 fallback, 자세 랜드마크 분석"),
        ("State/Data", "authStore, mock history", "데모 계정, 학습 이력, 설정 상태 저장"),
        ("UX System", "Button, AppDialog, Sidebar", "일관된 버튼, 알림/확인 모달, 내비게이션"),
    ], [1.1, 2.0, 3.4])
    doc.add_heading("5. 구현 결과", level=1)
    for title, desc in [
        ("브랜딩 전환", "PostureAI 명칭을 Moti로 전환하고 앱/설치 아이콘을 M 로고로 통일하였다."),
        ("로그인 편의 개선", "admin/admin 계정을 기본 제공하면서 회원가입 흐름은 유지하였다."),
        ("학습 화면 개선", "웹캠을 크게 표시하고, 카메라 상태를 상단에 배치하며, 하단 시간 표시와 불필요한 여백을 제거하였다."),
        ("이력 캘린더", "일간/주간/월간 보기, 기간 이동, 오늘 버튼, 주말 색상 표시를 구현하였다."),
        ("커스텀 알림", "기본 alert/confirm을 제거하고 UI 톤에 맞는 다이얼로그로 변경하였다."),
    ]:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)
    doc.add_heading("6. 테스트 및 검증", level=1)
    add_table(doc, ["검증 항목", "결과"], [
        ("TypeScript 타입체크", "tsc -b 통과"),
        ("웹캠 해상도 오류", "OverconstrainedError 발생 시 ideal 1280x720 fallback 적용"),
        ("학습 화면 UI", "최신 캡처 기준 웹캠 영역, 상단 상태, 컨트롤 정렬 확인"),
        ("산출물", "기능 명세서, 보고서, 기획서, 아키텍처, 스크린 flow 재생성"),
    ], [2.3, 4.2])
    doc.add_heading("7. 결론 및 향후 개선", level=1)
    doc.add_paragraph("Moti는 현재 데모/시연 가능한 수준의 데스크톱 자세 학습 앱으로 구성되었다. 향후 실제 백엔드, DB, AI 피드백 API, 장기 학습 데이터 분석, 카메라 장치 선택 기능을 연결하면 실사용 서비스로 확장할 수 있다.")
    doc.save(REPORTS / "Moti_최종보고서.docx")


def html_page(title: str, body: str) -> str:
    return f"""<!doctype html>
<html lang="ko">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{title}</title>
  <style>
    :root {{ --green:#58cc02; --blue:#1cb0f6; --red:#ff4b4b; --ink:#1f2937; --muted:#667085; --bg:#f6f8fa; --card:#fff; --line:#dde5ed; }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; background:var(--bg); color:var(--ink); font-family:'Malgun Gothic','Apple SD Gothic Neo',Arial,sans-serif; line-height:1.7; }}
    .cover {{ padding:72px 64px; color:white; background:linear-gradient(135deg,#1d1d1f,#19331d 55%,#17384d); }}
    .cover h1 {{ margin:0; font-size:48px; letter-spacing:-1px; }}
    .cover p {{ max-width:820px; color:#d7e6d9; font-size:18px; }}
    main {{ max-width:1120px; margin:32px auto 80px; padding:0 24px; }}
    section {{ background:var(--card); border:1px solid var(--line); border-radius:10px; padding:28px; margin-bottom:22px; box-shadow:0 12px 28px rgba(31,41,55,.06); }}
    h2 {{ margin:0 0 16px; font-size:26px; }}
    h3 {{ margin:24px 0 10px; font-size:18px; }}
    .grid {{ display:grid; grid-template-columns:repeat(3,minmax(0,1fr)); gap:14px; }}
    .card {{ border:1px solid var(--line); border-radius:8px; padding:16px; background:#fbfcfd; }}
    .tag {{ display:inline-flex; padding:4px 9px; border-radius:999px; background:#ecfbe5; color:#247000; font-weight:800; font-size:12px; }}
    table {{ width:100%; border-collapse:collapse; margin:12px 0; font-size:14px; }}
    th,td {{ border:1px solid var(--line); padding:10px 12px; vertical-align:top; }}
    th {{ background:#ecfbe5; text-align:left; }}
    code {{ background:#eef2f6; border-radius:5px; padding:2px 5px; }}
    .flow {{ width:100%; border:1px solid var(--line); border-radius:8px; }}
  </style>
</head>
<body>
  <div class="cover"><span class="tag">Moti</span><h1>{title}</h1><p>AI 기반 VDT 증후군 자세 교정 데스크톱 앱 산출물</p></div>
  <main>{body}</main>
</body>
</html>"""


def create_feature_html() -> None:
    rows = "\n".join(f"<tr><td>{fid}</td><td>{name}</td><td>{desc}</td><td>{status}</td><td>{module}</td></tr>" for fid, name, desc, status, module in FEATURES)
    body = f"""
    <section><h2>기능 명세서</h2><p>Moti의 현재 구현 기능을 기능 ID 단위로 정리했다.</p></section>
    <section><h2>전체 기능 목록</h2><table><thead><tr><th>ID</th><th>기능</th><th>상세 설명</th><th>상태</th><th>관련 화면/모듈</th></tr></thead><tbody>{rows}</tbody></table></section>
    """
    (REPORTS / "Moti_기능명세서.html").write_text(html_page("Moti 기능 명세서", body), encoding="utf-8")


def create_plan_html() -> None:
    body = """
    <section><h2>1. 프로젝트 개요</h2><p>Moti는 PC 웹캠을 활용해 사용자의 자세를 실시간으로 확인하고, 학습 모드별 자세 점수와 피드백을 제공하는 데스크톱 애플리케이션이다.</p></section>
    <section><h2>2. 기획 배경</h2><div class="grid"><div class="card"><h3>문제</h3><p>장시간 PC 사용으로 목, 어깨, 손목, 눈 피로가 누적된다.</p></div><div class="card"><h3>해결</h3><p>웹캠과 AI 자세 분석으로 자세 문제를 시각화한다.</p></div><div class="card"><h3>가치</h3><p>시연 가능한 데스크톱 UX와 학습 이력 기반 관리 경험을 제공한다.</p></div></div></section>
    <section><h2>3. 핵심 사용자 흐름</h2><ol><li>로그인 또는 회원가입</li><li>대시보드에서 학습 모드 선택</li><li>웹캠 화면 확인 후 자세교정 시작</li><li>학습 결과를 통계와 이력에서 확인</li><li>설정에서 알림, 테마, 개인정보 관련 기능 관리</li></ol></section>
    <section><h2>4. 주요 기능</h2><table><tr><th>영역</th><th>내용</th></tr><tr><td>인증</td><td>admin/admin 기본 로그인, 회원가입, 아이디/비밀번호 찾기</td></tr><tr><td>학습</td><td>거북목, 키보드, 어깨, 안구 모드와 웹캠 기반 자세 점검</td></tr><tr><td>통계</td><td>점수 변화, 위험도, AI 피드백, 추천 스트레칭</td></tr><tr><td>이력</td><td>일간/주간/월간 캘린더와 날짜별 세션 표시</td></tr></table></section>
    <section><h2>5. 개발 및 시연 범위</h2><p>현재 범위는 Electron/React 프론트엔드 중심의 시연 가능 버전이며, 향후 실제 백엔드, DB, AI 피드백 API 연동을 통해 서비스화할 수 있다.</p></section>
    """
    (REPORTS / "Moti_기획서.html").write_text(html_page("Moti 프로젝트 기획서", body), encoding="utf-8")


def create_architecture() -> None:
    md = """# Moti 시스템 아키텍처

## 1. 개요

Moti는 Electron 데스크톱 Shell 위에서 React Renderer를 실행하고, 브라우저 WebRTC API와 MediaPipe Pose를 이용해 웹캠 기반 자세 분석을 수행한다. 현재 저장소는 시연 가능한 프론트엔드 중심 구조이며, 인증/학습 이력은 localStorage와 mock data를 사용한다.

## 2. 전체 구조

```mermaid
flowchart TB
  User["사용자"] --> Electron["Electron Desktop App"]
  Camera["PC 웹캠"] --> WebRTC["navigator.mediaDevices.getUserMedia"]

  subgraph Electron["Electron Shell"]
    Main["electron/main.ts"]
    Preload["electron/preload.ts"]
    Renderer["React Renderer"]
  end

  subgraph Renderer["React App"]
    App["App.tsx / HashRouter"]
    Dialog["AppDialog 커스텀 알림"]
    Sidebar["Sidebar"]
    Auth["Login / Register / Find ID / Find Password"]
    Dashboard["Dashboard / ModeSelector"]
    Learning["LearningSession / PostureMonitor"]
    Stats["Statistics"]
    History["LearningHistory"]
    Settings["Settings"]
  end

  subgraph Vision["AI Vision Layer"]
    Webcam["useWebcam: 1280x720 exact -> ideal fallback"]
    MediaPipe["useMediaPipe: Pose 모델 로딩"]
    Canvas["Canvas 관절 표시"]
    Calculator["postureCalculator"]
  end

  subgraph LocalData["Local Data"]
    AuthStore["authStore / localStorage"]
    MockHistory["mockLearningHistory"]
    Modes["modes.tsx"]
  end

  App --> Dialog
  App --> Auth
  App --> Sidebar
  App --> Dashboard
  App --> Learning
  App --> Stats
  App --> History
  App --> Settings
  Dashboard --> Modes
  Learning --> Webcam
  Webcam --> WebRTC
  Learning --> MediaPipe
  MediaPipe --> Canvas
  MediaPipe --> Calculator
  Auth --> AuthStore
  History --> MockHistory
```

## 3. 화면 라우팅

| Route | 화면 | 역할 |
| --- | --- | --- |
| `/#/login` | 로그인 | admin/admin 로그인, 계정 찾기/회원가입 이동 |
| `/#/register` | 회원가입 | 신규 계정 등록 |
| `/#/find-id` | 아이디 찾기 | 계정 찾기 보조 흐름 |
| `/#/find-password` | 비밀번호 찾기 | 비밀번호 찾기 보조 흐름 |
| `/#/dashboard` | 대시보드 | 학습 모드 진입 |
| `/#/learn/:modeId` | 자세 학습 | 웹캠/AI 분석/관절 표시 |
| `/#/statistics` | 통계 | 점수 변화, 위험도, AI 피드백 |
| `/#/history` | 학습 이력 | 일/주/월 캘린더와 날짜별 기록 |
| `/#/settings` | 설정 | 알림, 테마, 데이터 관리 |

## 4. 학습 실행 흐름

```mermaid
sequenceDiagram
  participant U as 사용자
  participant UI as LearningSession
  participant Cam as useWebcam
  participant AI as useMediaPipe
  participant Canvas as Canvas Overlay

  U->>UI: 자세교정 시작 클릭
  UI->>Cam: getUserMedia(1280x720 exact)
  Cam-->>UI: 실패 시 ideal 1280x720 fallback
  UI->>AI: MediaPipe Pose 초기화
  AI->>AI: video frame 분석
  AI->>Canvas: 랜드마크/관절선 렌더링
  UI-->>U: 점수, 현재 상태, 미리보기 상태 표시
  U->>UI: 자세교정 중지 클릭
  UI->>AI: 분석 루프 중지
  UI->>Cam: track stop
```

## 5. 현재 구현 기준

- 앱명과 설치 아이콘은 Moti 기준으로 정리되어 있다.
- 설치 동의문은 Moti 개인정보/카메라 동의 내용으로 구성한다.
- custom dialog를 사용해 기본 alert/confirm을 대체한다.
- 학습 화면은 웹캠 확인 화면 우측에 카메라 미리보기 상태를 표시한다.
- 하단 시간초 표시를 제거하고 웹캠 영역을 크게 보여준다.
- 학습 이력은 일간/주간/월간 보기, 이전/다음/오늘 이동, 주말 색상을 지원한다.

## 6. 향후 확장

- 백엔드 API와 DB(users, learning_sessions, posture_logs, daily_statistics) 연결
- 실제 AI 피드백 API 연동
- 카메라 장치 선택 기능 실제 연결
- 학습 결과 저장 및 장기 추세 분석
- 배포 파이프라인과 자동화 테스트 구성
"""
    (REPORTS / "Moti_시스템_아키텍처.md").write_text(md, encoding="utf-8")
    body = "<section><h2>시스템 아키텍처</h2><p>아래 Mermaid 블록은 Markdown 파일에서 확인할 수 있는 전체 구조와 실행 흐름을 HTML 요약으로 정리한 것이다.</p></section>"
    body += "<section><h2>핵심 구조</h2><div class='grid'><div class='card'><h3>Electron</h3><p>데스크톱 앱 Shell, 창, 설치 패키징</p></div><div class='card'><h3>React</h3><p>화면 라우팅과 사용자 인터페이스</p></div><div class='card'><h3>MediaPipe</h3><p>웹캠 프레임 자세 랜드마크 분석</p></div></div></section>"
    body += "<section><h2>상세 문서</h2><p><code>Moti_시스템_아키텍처.md</code> 파일에 Mermaid 다이어그램과 라우팅 표를 포함했다.</p></section>"
    (REPORTS / "Moti_시스템_아키텍처.html").write_text(html_page("Moti 시스템 아키텍처", body), encoding="utf-8")


def img_data(file: str) -> str:
    path = SCREENSHOTS / file
    if not path.exists():
        return ""
    return "data:image/png;base64," + base64.b64encode(path.read_bytes()).decode("ascii")


def create_screen_flow() -> None:
    screens = [
        ("01-login.png", "로그인", "admin/admin 또는 회원가입/계정 찾기"),
        ("02-register.png", "회원가입", "신규 사용자 등록"),
        ("03-find-id.png", "아이디 찾기", "아이디 확인 보조 흐름"),
        ("04-find-password.png", "비밀번호 찾기", "비밀번호 찾기 보조 흐름"),
        ("05-dashboard.png", "대시보드", "학습 모드 선택"),
        ("06-learning-session.png", "자세 학습", "웹캠 미리보기와 AI 분석"),
        ("07-statistics.png", "통계", "점수 변화와 AI 피드백"),
        ("08-learning-history.png", "학습 이력", "일/주/월 캘린더"),
        ("09-settings.png", "설정", "알림/테마/데이터 관리"),
    ]
    cards = []
    positions = [(80,120),(80,520),(440,520),(440,850),(840,120),(1320,120),(1320,560),(840,560),(840,900)]
    for (file, title, desc), (x, y) in zip(screens, positions):
        data = img_data(file)
        cards.append(f"""
<g>
  <rect x="{x}" y="{y}" width="320" height="300" rx="10" fill="#fff" stroke="#dde5ed"/>
  <image x="{x}" y="{y}" width="320" height="210" href="{data}" preserveAspectRatio="xMidYMin slice"/>
  <text x="{x+16}" y="{y+245}" font-size="20" font-weight="800" fill="#1f2937">{title}</text>
  <text x="{x+16}" y="{y+272}" font-size="14" fill="#667085">{desc}</text>
</g>""")
    arrows = """
<defs><marker id="a" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3 z" fill="#58cc02"/></marker></defs>
<path d="M400 250 C560 250 680 250 840 250" stroke="#58cc02" stroke-width="4" fill="none" marker-end="url(#a)"/>
<path d="M1160 250 C1220 250 1260 250 1320 250" stroke="#58cc02" stroke-width="4" fill="none" marker-end="url(#a)"/>
<path d="M1480 420 C1480 480 1480 500 1480 560" stroke="#58cc02" stroke-width="4" fill="none" marker-end="url(#a)"/>
<path d="M1320 710 C1240 710 1230 710 1160 710" stroke="#58cc02" stroke-width="4" fill="none" marker-end="url(#a)"/>
<path d="M240 420 C240 470 240 480 240 520" stroke="#1cb0f6" stroke-width="3" fill="none" marker-end="url(#a)"/>
<path d="M400 330 C500 410 560 470 600 520" stroke="#1cb0f6" stroke-width="3" fill="none" marker-end="url(#a)"/>
"""
    svg = f"""<svg xmlns="http://www.w3.org/2000/svg" width="1720" height="1240" viewBox="0 0 1720 1240">
<rect width="1720" height="1240" fill="#f6f8fa"/>
<text x="70" y="70" font-size="36" font-weight="900" fill="#1f2937">Moti Screen Flow</text>
<text x="70" y="100" font-size="16" fill="#667085">로그인부터 학습, 이력, 통계, 설정까지 이어지는 화면 흐름</text>
{arrows}
{''.join(cards)}
</svg>"""
    (REPORTS / "Moti_스크린_Flow.svg").write_text(svg, encoding="utf-8")

    html_cards = "\n".join(
        f"<article class='card'><img src='screenshots/{file}'><h3>{title}</h3><p>{desc}</p></article>"
        for file, title, desc in screens
    )
    body = f"""
    <section><h2>스크린 Flow</h2><p>Moti의 주요 화면 흐름을 한 화면에서 설명하기 위한 자료다.</p><img class="flow" src="Moti_스크린_Flow.svg" alt="Moti screen flow"></section>
    <section><h2>화면 목록</h2><div class="grid">{html_cards}</div></section>
    <style>.card img{{width:100%;height:170px;object-fit:cover;object-position:top;border-radius:6px;border:1px solid var(--line);}}</style>
    """
    (REPORTS / "Moti_스크린_Flow.html").write_text(html_page("Moti 스크린 Flow", body), encoding="utf-8")


def main() -> None:
    REPORTS.mkdir(exist_ok=True)
    create_feature_spec()
    create_feature_html()
    create_final_report()
    create_plan_html()
    create_architecture()
    create_screen_flow()


if __name__ == "__main__":
    main()
